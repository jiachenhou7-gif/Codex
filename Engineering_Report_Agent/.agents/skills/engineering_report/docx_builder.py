import json
import os
from docx import Document
from docx.shared import RGBColor

def apply_style(paragraph, style_name):
    try:
        paragraph.style = style_name
    except KeyError:
        fallback = "Normal"
        try:
            paragraph.style = fallback
        except KeyError:
            paragraph.style = "Normal"
        print(f"Warning: style not found, using '{paragraph.style.name}'.")

def build_word_by_styles():
    json_path = "temp_report.json"
    template_path = "templates/report_template.docx" # 你的样式模板文档
    output_path = "output_report.docx"

    if not os.path.exists(json_path):
        print("错误：未找到生成的中间数据 temp_report.json")
        return

    # 1. 读取 Agent 写入的结构化文本块
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # 2. 加载已经内置好所有样式的空白 Word 模板
    doc = Document(template_path)
    
    # 清理空白模板可能自带的第一个空白段落（可选）
    if len(doc.paragraphs) == 1 and doc.paragraphs[0].text == "":
        p_first = doc.paragraphs[0]
    else:
        p_first = None

    # 3. 按顺序循环写入内容块
    for index, block in enumerate(data.get("blocks", [])):
        block_type = block.get("type", "text") # text 或 table
        style_name = block.get("style", "正文") # 对应 Word 中的样式名
        source = block.get("source", "normal")
        
        # --- 处理文本内容（标题、正文、图名、表名） ---
        if block_type == "text":
            # 如果是第一条数据且文档首段为空，直接复用首段，否则追加新段落
            if index == 0 and p_first:
                p = p_first
            else:
                p = doc.add_paragraph()
            
            apply_style(p, style_name)  # 直接调用 Word 中对应的样式名

            # 联网数据特殊染色处理
            if source == "web":
                # 插入红色警告前缀
                warning_run = p.add_run("【⚠️ 联网数据审阅】")
                warning_run.font.color.rgb = RGBColor(255, 0, 0)
                warning_run.bold = True
                
                # 网络正文染成深蓝色
                body_run = p.add_run(block["text"])
                body_run.font.color.rgb = RGBColor(0, 102, 204)
            else:
                p.add_run(block["text"])

        # --- 处理数据表格（如土石方平衡表） ---
        elif block_type == "table":
            table_data = block.get("data", [])
            if not table_data:
                continue
                
            rows = len(table_data)
            cols = len(table_data[0])
            
            # 创建表格
            table = doc.add_table(rows=rows, cols=cols)
            # 应用 Word 自带的无边框/普通表格样式，排版由单元格内段落样式决定
            table.style = 'Table Grid' 
            
            for r_idx, row_content in enumerate(table_data):
                for c_idx, cell_value in enumerate(row_content):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = str(cell_value)
                    
                    # 强制将表格内文字的段落样式指定为“表格内容”
                    p_cell = cell.paragraphs[0]
                    apply_style(p_cell, '表格内容')

    # 4. 保存为最终成品报告
    doc.save(output_path)
    print(f"Success: report generated at {output_path}")

if __name__ == "__main__":
    build_word_by_styles()
